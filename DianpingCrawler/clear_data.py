#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
__author__ = 'AJay'
__mtime__ = '2019/3/21 0021'

"""
import requests
from docx import Document
from docx.shared import Inches,Pt
from docx.oxml.ns import qn
key_word =['小孩','儿童','儿子','女儿','小孩子','娃','嬉戏','亲子','玩具','宝宝','宝贝']
childer =[]
with open('code_dict.txt','r',encoding='utf-8')as f:
    content = f.readlines()
    for one_line in content:
        dict = eval(one_line)

        for w in key_word:
            if w in dict['comment']:
                childer.append(dict)
                break

print(len(childer))
print((childer[1:10]))


doc =Document()

for i in childer:
    print(i)
    name = i.get('name')
    comment = i.get('comment')
    time = i.get('time')
    star = i.get('star')
    pic = i.get('pic')

    pen = doc.add_paragraph()
    ph = pen.paragraph_format
    ph.line_spacing = Pt(22)
    pensize1 = pen.add_run('用户：'+name+'\n')
    pensize = pen.add_run('评分：'+str(star)+'\n')
    pensize2 = pen.add_run('时间：'+time+'\n')
    pensize3 = pen.add_run('评论：'+comment+'\n')

    pensize.font.name = '宋体'
    pensize._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pensize.font.size = Pt(15)
    pensize.bold=True

    pensize1.font.name = '宋体'
    pensize1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pensize1.font.size = Pt(15)

    pensize2.font.name = '宋体'
    pensize2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pensize2.font.size = Pt(15)

    pensize3.font.name = '宋体'
    pensize3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pensize3.font.size = Pt(15)
    if pic:
        for p in pic:
            req = requests.get(p)
            with open('capth.png','wb')as f:
                f.write(req.content)

            doc.add_picture('capth.png', width=Inches(2.5))

doc.save('dianping.docx')